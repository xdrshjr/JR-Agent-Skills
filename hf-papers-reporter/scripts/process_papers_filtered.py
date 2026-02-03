#!/usr/bin/env python3
"""
Hugging Face Daily Papers 报告生成器 - 图片筛选改进版
"""

import requests
import json
import os
import re
from pathlib import Path
from bs4 import BeautifulSoup
import fitz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PIL import Image
import numpy as np

# 配置
WORK_DIR = Path(__file__).parent.parent  # skill根目录
PDF_DIR = WORK_DIR / "pdfs"
IMAGE_DIR = WORK_DIR / "images"
OUTPUT_DIR = WORK_DIR / "output"

PAPERS = [
    {"title": "Idea2Story: An Automated Pipeline for Transforming Research Concepts into Complete Scientific Narratives", "arxiv_id": "2601.20833", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.20833.png"},
    {"title": "Everything in Its Place: Benchmarking Spatial Intelligence of Text-to-Image Models", "arxiv_id": "2601.20354", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.20354.png"},
    {"title": "Scaling Embeddings Outperforms Scaling Experts in Language Models", "arxiv_id": "2601.21204", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21204.png"},
    {"title": "DynamicVLA: A Vision-Language-Action Model for Dynamic Object Manipulation", "arxiv_id": "2601.22153", "img_url": "https://cdn-avatars.huggingface.co/v1/production/uploads/63f47b5321eb234ab739e91a/vWfFNVtMkHl8gieha5PPd.jpeg"},
    {"title": "OCRVerse: Towards Holistic OCR in End-to-End Vision-Language Models", "arxiv_id": "2601.21639", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21639.png"},
    {"title": "MMFineReason: Closing the Multimodal Reasoning Gap via Open Data-Centric Methods", "arxiv_id": "2601.21821", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21821.png"},
    {"title": "ConceptMoE: Adaptive Token-to-Concept Compression for Implicit Compute Allocation", "arxiv_id": "2601.21420", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21420.png"},
    {"title": "PLANING: A Loosely Coupled Triangle-Gaussian Framework for Streaming 3D Reconstruction", "arxiv_id": "2601.22046", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.22046.png"},
    {"title": "Qwen3-ASR Technical Report", "arxiv_id": "2601.21337", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21337.png"},
    {"title": "AgentLongBench: A Controllable Long Benchmark For Long-Contexts Agents via Environment Rollouts", "arxiv_id": "2601.20730", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.20730.png"},
]

HEADERS = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}

# 图片筛选参数
MIN_IMAGE_WIDTH = 150      # 最小宽度（过滤小图标）
MIN_IMAGE_HEIGHT = 100     # 最小高度
MAX_IMAGE_WIDTH = 2000     # 最大宽度（过滤异常大图）
MAX_IMAGE_HEIGHT = 1500    # 最大高度
MIN_CONTENT_RATIO = 0.05   # 最小内容比例（过滤空白/纯背景图片）
MAX_TEXT_LIKE_RATIO = 0.85 # 最大类文本比例（过滤纯文本截图）

def download_pdf(arxiv_id):
    """从 arxiv.org 下载 PDF"""
    pdf_path = PDF_DIR / f"{arxiv_id}.pdf"
    if pdf_path.exists():
        with open(pdf_path, 'rb') as f:
            header = f.read(10)
            if header.startswith(b'%PDF'):
                return str(pdf_path)
    
    url = f"https://arxiv.org/pdf/{arxiv_id}.pdf"
    try:
        resp = requests.get(url, headers=HEADERS, timeout=60)
        if resp.status_code == 200 and resp.content.startswith(b'%PDF'):
            with open(pdf_path, 'wb') as f:
                f.write(resp.content)
            return str(pdf_path)
        else:
            print(f"    下载失败或不是PDF: {resp.status_code}")
    except Exception as e:
        print(f"    下载错误: {e}")
    return None

def download_cover_image(img_url, arxiv_id):
    """下载并压缩封面图"""
    img_path = IMAGE_DIR / f"{arxiv_id}_cover.png"
    if img_path.exists():
        return str(img_path)
    
    try:
        resp = requests.get(img_url, headers=HEADERS, timeout=30)
        if resp.status_code == 200:
            img = Image.open(BytesIO(resp.content))
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            img.thumbnail((800, 600))
            img.save(img_path, "PNG")
            return str(img_path)
    except Exception as e:
        print(f"    封面图错误: {e}")
    return None

def is_likely_figure(img_data):
    """
    判断图片是否可能是论文中的插图/图表
    返回: (是否通过, 原因)
    """
    width, height = img_data.size
    
    # 1. 尺寸检查
    if width < MIN_IMAGE_WIDTH or height < MIN_IMAGE_HEIGHT:
        return False, f"太小 ({width}x{height})"
    
    if width > MAX_IMAGE_WIDTH or height > MAX_IMAGE_HEIGHT:
        return False, f"太大 ({width}x{height})"
    
    # 2. 长宽比检查（过滤极端比例）
    aspect_ratio = width / height
    if aspect_ratio < 0.3 or aspect_ratio > 5:
        return False, f"极端比例 ({aspect_ratio:.2f})"
    
    # 3. 内容分析
    img_array = np.array(img_data)
    
    # 转换为灰度分析
    if len(img_array.shape) == 3:
        gray = np.mean(img_array, axis=2)
    else:
        gray = img_array
    
    # 计算内容比例（非纯白/纯黑像素）
    total_pixels = gray.size
    non_blank = np.sum((gray < 250) & (gray > 10))
    content_ratio = non_blank / total_pixels
    
    if content_ratio < MIN_CONTENT_RATIO:
        return False, f"内容太少 ({content_ratio:.2%})"
    
    # 4. 颜色多样性检查（过滤单色/双色图片，可能是页眉页脚）
    if len(img_array.shape) == 3:
        # 下采样后统计颜色
        small = img_data.resize((50, 50))
        small_array = np.array(small)
        unique_colors = len(np.unique(small_array.reshape(-1, 3), axis=0))
        color_ratio = unique_colors / (50 * 50)
        
        # 如果颜色种类太少，可能是简单图标或页眉
        if color_ratio < 0.05 and content_ratio < 0.3:
            return False, f"颜色太单一 ({color_ratio:.2%})"
    
    # 5. 检查是否像纯文本（通过边缘检测原理）
    # 论文插图通常有图形、线条、颜色区域
    # 纯文本图片通常有大量水平条纹
    grad_x = np.abs(np.diff(gray, axis=1, append=gray[:, -1:]))
    grad_y = np.abs(np.diff(gray, axis=0, append=gray[-1:, :]))
    
    # 如果水平梯度远大于垂直梯度，可能是纯文本
    if np.mean(grad_x) > 0 and np.mean(grad_y) / np.mean(grad_x) < 0.1:
        # 可能是纯文本行
        if content_ratio > 0.7:  # 文本通常填充率高
            return False, "可能是纯文本"
    
    return True, "通过"

def extract_from_pdf(pdf_path):
    """从PDF提取文本和图片（带筛选）"""
    result = {"abstract": "", "introduction": "", "images": []}
    
    try:
        doc = fitz.open(pdf_path)
        
        # 提取文本（前5页）
        full_text = ""
        for page_num in range(min(5, len(doc))):
            full_text += doc[page_num].get_text()
        
        # 提取 Abstract
        abstract_patterns = [
            r'Abstract[.\s]*(.+?)(?=\n\s*\n|Introduction|\d+\s+Introduction|I\s+INTRODUCTION)',
            r'ABSTRACT[.\s]*(.+?)(?=\n\s*\n|INTRODUCTION)',
        ]
        for pattern in abstract_patterns:
            match = re.search(pattern, full_text, re.DOTALL | re.IGNORECASE)
            if match:
                result["abstract"] = match.group(1).strip()[:3000]
                break
        
        # 提取 Introduction
        intro_patterns = [
            r'(?:Introduction|1\s+Introduction)[.\s]*(.+?)(?=\n\s*\d+\s+\w|Related Work|Method|Approach|Background|Preliminaries|\d+\s+RELATED)',
            r'I\s+INTRODUCTION[.\s]*(.+?)(?=II\s+\w|RELATED|METHOD)',
        ]
        for pattern in intro_patterns:
            match = re.search(pattern, full_text, re.DOTALL | re.IGNORECASE)
            if match:
                result["introduction"] = match.group(1).strip()[:8000]
                break
        
        # 提取图片（带筛选）
        img_count = 0
        filtered_count = 0
        
        for page_num in range(min(5, len(doc))):
            page = doc[page_num]
            images = page.get_images()
            
            for img_idx, img in enumerate(images):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # 处理图片
                    img_data = Image.open(BytesIO(image_bytes))
                    if img_data.mode in ('RGBA', 'P'):
                        img_data = img_data.convert('RGB')
                    
                    # 筛选图片
                    is_valid, reason = is_likely_figure(img_data)
                    
                    if not is_valid:
                        filtered_count += 1
                        print(f"    过滤图片 {img_idx}: {reason}")
                        continue
                    
                    # 压缩并保存
                    img_data.thumbnail((600, 400))
                    img_path = IMAGE_DIR / f"{Path(pdf_path).stem}_img_{img_count}_filtered.png"
                    img_data.save(img_path, "PNG")
                    result["images"].append(str(img_path))
                    img_count += 1
                    
                    # 每篇论文最多4张图
                    if img_count >= 4:
                        break
                        
                except Exception as e:
                    pass
            
            if img_count >= 4:
                break
        
        if filtered_count > 0:
            print(f"    过滤了 {filtered_count} 张不相关图片")
        
        doc.close()
    except Exception as e:
        print(f"    PDF处理错误: {e}")
    
    return result

def process_paper(paper_info, index):
    """处理单篇论文"""
    print(f"\n[{index+1}/10] {paper_info['title'][:50]}...")
    
    arxiv_id = paper_info['arxiv_id']
    result = {
        "title": paper_info['title'],
        "arxiv_id": arxiv_id,
        "cover_image": None,
        "pdf_images": [],
        "abstract": "",
        "introduction": ""
    }
    
    # 1. 下载 PDF
    pdf_path = download_pdf(arxiv_id)
    if not pdf_path:
        print(f"    ❌ PDF下载失败")
        return result
    print(f"    ✓ PDF下载完成")
    
    # 2. 下载封面图
    cover_path = download_cover_image(paper_info['img_url'], arxiv_id)
    if cover_path:
        result['cover_image'] = cover_path
        print(f"    ✓ 封面图")
    
    # 3. 提取内容（带图片筛选）
    content = extract_from_pdf(pdf_path)
    result['abstract'] = content['abstract']
    result['introduction'] = content['introduction']
    result['pdf_images'] = content['images']
    
    print(f"    ✓ 摘要: {len(content['abstract'])} 字符")
    print(f"    ✓ 介绍: {len(content['introduction'])} 字符")
    print(f"    ✓ 相关图片: {len(content['images'])} 张")
    
    return result

def create_word_report(results):
    """生成 Word 报告"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('Hugging Face Daily Papers 报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('生成日期: 2026年2月1日').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('（已过滤不相关图片）').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for idx, paper in enumerate(results):
        if not paper['abstract']:
            continue
        
        # 论文标题
        doc.add_heading(f"{idx+1}. {paper['title']}", level=1)
        doc.add_paragraph(f"arXiv: https://arxiv.org/abs/{paper['arxiv_id']}")
        
        # 封面图
        if paper['cover_image'] and os.path.exists(paper['cover_image']):
            try:
                doc.add_heading('封面图', level=2)
                doc.add_picture(paper['cover_image'], width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
        
        # 清理XML非法字符
        def clean_xml(text):
            if not text:
                return ""
            text = text.replace('\n', ' ').replace('\r', ' ').strip()
            cleaned = []
            for char in text:
                code = ord(char)
                if code < 0x20 and code not in (0x09, 0x0A, 0x0D):
                    continue
                cleaned.append(char)
            return ''.join(cleaned)
        
        # 摘要
        if paper['abstract']:
            doc.add_heading('摘要 (Abstract)', level=2)
            abstract = clean_xml(paper['abstract'])
            if abstract:
                doc.add_paragraph(abstract)
        
        # 介绍
        if paper['introduction']:
            doc.add_heading('介绍 (Introduction)', level=2)
            intro = clean_xml(paper['introduction'])
            if intro:
                doc.add_paragraph(intro)
        
        # 论文插图（已筛选）
        if paper['pdf_images']:
            doc.add_heading('论文插图', level=2)
            for img_path in paper['pdf_images']:
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                    except:
                        pass
        
        # 分隔
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
    
    output_path = OUTPUT_DIR / "HF_Daily_Papers_Report.docx"
    doc.save(output_path)
    return str(output_path)

def main():
    print("="*60)
    print("Hugging Face Daily Papers 报告生成器（图片筛选版）")
    print("="*60)
    
    # 清理旧的图片文件（保留封面图）
    for img in IMAGE_DIR.glob("*_img_*.png"):
        img.unlink()
    print("已清理旧图片文件\n")
    
    results = []
    for idx, paper in enumerate(PAPERS):
        result = process_paper(paper, idx)
        results.append(result)
    
    print("\n" + "="*60)
    print("生成 Word 报告...")
    report_path = create_word_report(results)
    print(f"✓ 报告已保存: {report_path}")
    
    # 保存JSON
    json_path = OUTPUT_DIR / "papers_data.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"✓ 数据已保存: {json_path}")
    
    # 统计
    success = sum(1 for r in results if r['abstract'])
    total_images = sum(len(r['pdf_images']) for r in results)
    print(f"\n成功处理: {success}/{len(PAPERS)} 篇论文")
    print(f"总共提取相关图片: {total_images} 张")

if __name__ == "__main__":
    main()
