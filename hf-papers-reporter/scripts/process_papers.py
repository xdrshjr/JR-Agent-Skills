#!/usr/bin/env python3
"""
Hugging Face Daily Papers 报告生成器 - 修订版
"""

import requests
import json
import os
import re
from pathlib import Path
from bs4 import BeautifulSoup
import fitz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PIL import Image

# 配置
WORK_DIR = Path(__file__).parent.parent  # skill根目录
PDF_DIR = WORK_DIR / "pdfs"
IMAGE_DIR = WORK_DIR / "images"
OUTPUT_DIR = WORK_DIR / "output"

PAPERS = [
    {"title": "Idea2Story: An Automated Pipeline for Transforming Research Concepts into Complete Scientific Narratives", "arxiv_id": "2601.20833", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.20833.png"},
    {"title": "Everything in Its Place: Benchmarking Spatial Intelligence of Text-to-Image Models", "arxiv_id": "2601.20354", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.20354.png"},
    {"title": "Scaling Embeddings Outperforms Scaling Experts in Language Models", "arxiv_id": "2601.21204", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21204.png"},
    {"title": "DynamicVLA: A Vision-Language-Action Model for Dynamic Object Manipulation", "arxiv_id": "2601.22153", "img_url": "https://cdn-avatars.huggingface.co/v1/production/uploads/63f47b5321eb234ab739e91a/vWfFNVtMkHl8gieha5PPd.jpeg"},
    {"title": "OCRVerse: Towards Holistic OCR in End-to-End Vision-Language Models", "arxiv_id": "2601.21639", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21639.png"},
    {"title": "MMFineReason: Closing the Multimodal Reasoning Gap via Open Data-Centric Methods", "arxiv_id": "2601.21821", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21821.png"},
    {"title": "ConceptMoE: Adaptive Token-to-Concept Compression for Implicit Compute Allocation", "arxiv_id": "2601.21420", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21420.png"},
    {"title": "PLANING: A Loosely Coupled Triangle-Gaussian Framework for Streaming 3D Reconstruction", "arxiv_id": "2601.22046", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.22046.png"},
    {"title": "Qwen3-ASR Technical Report", "arxiv_id": "2601.21337", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.21337.png"},
    {"title": "AgentLongBench: A Controllable Long Benchmark For Long-Contexts Agents via Environment Rollouts", "arxiv_id": "2601.20730", "img_url": "https://cdn-thumbnails.huggingface.co/social-thumbnails/papers/2601.20730.png"},
]

HEADERS = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}

def download_pdf(arxiv_id):
    """从 arxiv.org 下载 PDF"""
    pdf_path = PDF_DIR / f"{arxiv_id}.pdf"
    if pdf_path.exists():
        # 检查是否真的是PDF
        with open(pdf_path, 'rb') as f:
            header = f.read(10)
            if header.startswith(b'%PDF'):
                return str(pdf_path)
    
    # 正确的 arXiv PDF URL
    url = f"https://arxiv.org/pdf/{arxiv_id}.pdf"
    try:
        resp = requests.get(url, headers=HEADERS, timeout=60)
        if resp.status_code == 200 and resp.content.startswith(b'%PDF'):
            with open(pdf_path, 'wb') as f:
                f.write(resp.content)
            return str(pdf_path)
        else:
            print(f"    下载失败或不是PDF: {resp.status_code}")
    except Exception as e:
        print(f"    下载错误: {e}")
    return None

def download_cover_image(img_url, arxiv_id):
    """下载并压缩封面图"""
    img_path = IMAGE_DIR / f"{arxiv_id}_cover.png"
    if img_path.exists():
        return str(img_path)
    
    try:
        resp = requests.get(img_url, headers=HEADERS, timeout=30)
        if resp.status_code == 200:
            img = Image.open(BytesIO(resp.content))
            # 转换为RGB（防止RGBA问题）
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            img.thumbnail((800, 600))
            img.save(img_path, "PNG")
            return str(img_path)
    except Exception as e:
        print(f"    封面图错误: {e}")
    return None

def extract_from_pdf(pdf_path):
    """从PDF提取文本和图片"""
    result = {"abstract": "", "introduction": "", "images": []}
    
    try:
        doc = fitz.open(pdf_path)
        
        # 提取文本（前5页）
        full_text = ""
        for page_num in range(min(5, len(doc))):
            full_text += doc[page_num].get_text()
        
        # 提取 Abstract
        abstract_patterns = [
            r'Abstract[.\s]*(.+?)(?=\n\s*\n|Introduction|\d+\s+Introduction|I\s+INTRODUCTION)',
            r'ABSTRACT[.\s]*(.+?)(?=\n\s*\n|INTRODUCTION)',
        ]
        for pattern in abstract_patterns:
            match = re.search(pattern, full_text, re.DOTALL | re.IGNORECASE)
            if match:
                result["abstract"] = match.group(1).strip()[:3000]  # 限制长度
                break
        
        # 提取 Introduction
        intro_patterns = [
            r'(?:Introduction|1\s+Introduction)[.\s]*(.+?)(?=\n\s*\d+\s+\w|Related Work|Method|Approach|Background|Preliminaries|\d+\s+RELATED)',
            r'I\s+INTRODUCTION[.\s]*(.+?)(?=II\s+\w|RELATED|METHOD)',
        ]
        for pattern in intro_patterns:
            match = re.search(pattern, full_text, re.DOTALL | re.IGNORECASE)
            if match:
                result["introduction"] = match.group(1).strip()[:8000]  # 限制长度
                break
        
        # 提取图片（前5页，每页最多3张）
        img_count = 0
        for page_num in range(min(5, len(doc))):
            page = doc[page_num]
            images = page.get_images()
            
            for img_idx, img in enumerate(images[:3]):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # 处理图片
                    img_data = Image.open(BytesIO(image_bytes))
                    if img_data.mode in ('RGBA', 'P'):
                        img_data = img_data.convert('RGB')
                    img_data.thumbnail((600, 400))
                    
                    img_path = IMAGE_DIR / f"{Path(pdf_path).stem}_img_{img_count}.png"
                    img_data.save(img_path, "PNG")
                    result["images"].append(str(img_path))
                    img_count += 1
                except Exception as e:
                    pass
        
        doc.close()
    except Exception as e:
        print(f"    PDF处理错误: {e}")
    
    return result

def process_paper(paper_info, index):
    """处理单篇论文"""
    print(f"\n[{index+1}/10] {paper_info['title'][:50]}...")
    
    arxiv_id = paper_info['arxiv_id']
    result = {
        "title": paper_info['title'],
        "arxiv_id": arxiv_id,
        "cover_image": None,
        "pdf_images": [],
        "abstract": "",
        "introduction": ""
    }
    
    # 1. 下载 PDF
    pdf_path = download_pdf(arxiv_id)
    if not pdf_path:
        print(f"    ❌ PDF下载失败")
        return result
    print(f"    ✓ PDF下载完成")
    
    # 2. 下载封面图
    cover_path = download_cover_image(paper_info['img_url'], arxiv_id)
    if cover_path:
        result['cover_image'] = cover_path
        print(f"    ✓ 封面图")
    
    # 3. 提取内容
    content = extract_from_pdf(pdf_path)
    result['abstract'] = content['abstract']
    result['introduction'] = content['introduction']
    result['pdf_images'] = content['images']
    
    print(f"    ✓ 摘要: {len(content['abstract'])} 字符")
    print(f"    ✓ 介绍: {len(content['introduction'])} 字符")
    print(f"    ✓ 图片: {len(content['images'])} 张")
    
    return result

def create_word_report(results):
    """生成中文 Word 报告"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('Hugging Face Daily Papers 报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('生成日期: 2026年1月31日').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for idx, paper in enumerate(results):
        if not paper['abstract']:
            continue
        
        # 论文标题
        doc.add_heading(f"{idx+1}. {paper['title']}", level=1)
        doc.add_paragraph(f"arXiv: https://arxiv.org/abs/{paper['arxiv_id']}")
        
        # 封面图
        if paper['cover_image'] and os.path.exists(paper['cover_image']):
            try:
                doc.add_heading('封面图', level=2)
                doc.add_picture(paper['cover_image'], width=Inches(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
        
        # 清理XML非法字符的函数
        def clean_xml(text):
            if not text:
                return ""
            # 替换换行和清理控制字符
            text = text.replace('\n', ' ').replace('\r', ' ').strip()
            # 移除XML不允许的控制字符 (0x00-0x08, 0x0B-0x0C, 0x0E-0x1F)
            cleaned = []
            for char in text:
                code = ord(char)
                if code < 0x20 and code not in (0x09, 0x0A, 0x0D):
                    continue
                cleaned.append(char)
            return ''.join(cleaned)
        
        # 摘要
        if paper['abstract']:
            doc.add_heading('摘要 (Abstract)', level=2)
            abstract = clean_xml(paper['abstract'])
            if abstract:
                doc.add_paragraph(abstract)
        
        # 介绍
        if paper['introduction']:
            doc.add_heading('介绍 (Introduction)', level=2)
            intro = clean_xml(paper['introduction'])
            if intro:
                doc.add_paragraph(intro)
        
        # 论文插图
        if paper['pdf_images']:
            doc.add_heading('论文插图', level=2)
            for img_path in paper['pdf_images'][:4]:
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                    except:
                        pass
        
        # 分隔
        doc.add_paragraph("_" * 60)
        doc.add_paragraph()
    
    output_path = OUTPUT_DIR / "HF_Daily_Papers_Report.docx"
    doc.save(output_path)
    return str(output_path)

def main():
    print("="*60)
    print("Hugging Face Daily Papers 报告生成器")
    print("="*60)
    
    # 清理旧的错误文件
    for pdf in PDF_DIR.glob("*.pdf"):
        with open(pdf, 'rb') as f:
            if not f.read(4) == b'%PDF':
                pdf.unlink()
                print(f"已删除无效文件: {pdf.name}")
    
    results = []
    for idx, paper in enumerate(PAPERS):
        result = process_paper(paper, idx)
        results.append(result)
    
    print("\n" + "="*60)
    print("生成 Word 报告...")
    report_path = create_word_report(results)
    print(f"✓ 报告已保存: {report_path}")
    
    # 保存JSON
    json_path = OUTPUT_DIR / "papers_data.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"✓ 数据已保存: {json_path}")
    
    # 统计
    success = sum(1 for r in results if r['abstract'])
    print(f"\n成功处理: {success}/{len(PAPERS)} 篇论文")

if __name__ == "__main__":
    main()
