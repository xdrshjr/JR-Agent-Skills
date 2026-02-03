#!/usr/bin/env python3
"""
将Markdown转换为专业Word文档，支持中文字体设置
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import subprocess

def set_chinese_font(run, font_name):
    """设置中文字体"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_page_number(section):
    """添加页码到页脚"""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加页码字段
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    return run

def convert_md_to_word(md_file, output_file, title_font="黑体", body_font="仿宋"):
    """转换Markdown到Word"""
    
    # 1. 使用pandoc生成基础Word文档
    temp_docx = "/tmp/temp_report.docx"
    subprocess.run([
        "pandoc", md_file, "-o", temp_docx,
        "--toc", "--toc-depth=2"
    ], check=True)
    
    # 2. 打开并调整字体
    doc = Document(temp_docx)
    
    # 设置正文默认字体
    style = doc.styles['Normal']
    style.font.name = body_font
    style.font.size = Pt(12)
    
    # 设置标题样式
    for i in range(1, 4):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = title_font
            heading_style.font.bold = True
            if i == 1:
                heading_style.font.size = Pt(18)
            elif i == 2:
                heading_style.font.size = Pt(16)
            else:
                heading_style.font.size = Pt(14)
        except:
            pass
    
    # 应用字体到所有段落
    for para in doc.paragraphs:
        is_heading = para.style.name.startswith('Heading')
        font_to_use = title_font if is_heading else body_font
        
        for run in para.runs:
            set_chinese_font(run, font_to_use)
    
    # 3. 添加页眉页脚
    for section in doc.sections:
        # 页眉
        header = section.header
        if header.paragraphs:
            header_para = header.paragraphs[0]
            header_para.text = "技术分析报告"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_para.runs:
                set_chinese_font(run, body_font)
                run.font.size = Pt(10)
        
        # 页脚（页码）
        page_num_run = add_page_number(section)
        set_chinese_font(page_num_run, body_font)
        page_num_run.font.size = Pt(10)
    
    # 4. 保存
    doc.save(output_file)
    print(f"✅ Word文档已生成: {output_file}")

def main():
    if len(sys.argv) < 3:
        print("Usage: python convert_to_word.py <input.md> <output.docx> [title_font] [body_font]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2]
    title_font = sys.argv[3] if len(sys.argv) > 3 else "黑体"
    body_font = sys.argv[4] if len(sys.argv) > 4 else "仿宋"
    
    convert_md_to_word(md_file, output_file, title_font, body_font)

if __name__ == "__main__":
    main()
