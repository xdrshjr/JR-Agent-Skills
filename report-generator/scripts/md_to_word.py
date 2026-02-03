#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Converter
将构成主义报告转换为Word文档
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for key in ["sz", "val", "color"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcPr.append(element)

def create_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    return doc

def add_heading_custom(doc, text, level):
    """添加自定义标题"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(50, 50, 50)
        else:
            run.font.size = Pt(12)
    return heading

def add_paragraph_custom(doc, text, style=None, bold=False, italic=False):
    """添加自定义段落"""
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return p

def process_markdown_table(doc, lines, start_idx):
    """处理Markdown表格"""
    table_lines = []
    i = start_idx
    
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i])
        i += 1
    
    if len(table_lines) < 3:
        return i
    
    # 解析表格
    rows = []
    for line in table_lines:
        if '---' in line:  # 跳过分隔行
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return i
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell.text = cell_text
                # 设置单元格字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                        run.font.size = Pt(10)
                # 第一行加粗
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()  # 表格后添加空行
    return i

def process_code_block(doc, lines, start_idx):
    """处理代码块"""
    code_lines = []
    i = start_idx + 1  # 跳过开头的 ```
    
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i])
        i += 1
    
    # 添加代码块（使用等宽字体）
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)
    
    # 添加浅灰色背景（通过底纹）
    p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_paragraph()  # 代码块后添加空行
    return i + 1

def process_inline_formatting(paragraph, text):
    """处理行内格式（粗体、斜体、行内代码）"""
    # 处理行内代码 `code`
    parts = re.split(r'`([^`]+)`', text)
    for i, part in enumerate(parts):
        if i % 2 == 1:  # 代码部分
            run = paragraph.add_run(part)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 50, 50)
        else:  # 普通文本
            # 处理粗体 **text**
            bold_parts = re.split(r'\*\*([^*]+)\*\*', part)
            for j, bold_part in enumerate(bold_parts):
                if j % 2 == 1:  # 粗体部分
                    run = paragraph.add_run(bold_part)
                    run.bold = True
                else:  # 处理斜体 *text*
                    italic_parts = re.split(r'\*([^*]+)\*', bold_part)
                    for k, italic_part in enumerate(italic_parts):
                        if k % 2 == 1:  # 斜体部分
                            run = paragraph.add_run(italic_part)
                            run.italic = True
                        else:
                            if italic_part:
                                run = paragraph.add_run(italic_part)

def process_image_reference(doc, line, images_dir):
    """处理图片引用"""
    match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
    if match:
        alt_text = match.group(1)
        img_path = match.group(2)
        
        # 尝试多种路径
        possible_paths = [
            Path(img_path),
            Path(images_dir) / Path(img_path).name,
            Path(images_dir) / img_path.replace('images/', ''),
        ]
        
        for path in possible_paths:
            if isinstance(path, Path) and path.exists():
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(str(path), width=Inches(5))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 添加图注
                    if alt_text:
                        caption = doc.add_paragraph()
                        caption_run = caption.add_run(alt_text)
                        caption_run.font.name = 'Microsoft YaHei'
                        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                        caption_run.font.size = Pt(9)
                        caption_run.italic = True
                        caption_run.font.color.rgb = RGBColor(100, 100, 100)
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    return True
                except Exception as e:
                    print(f"  图片添加失败: {e}")
                    break
        
        # 图片不存在，添加占位符
        p = doc.add_paragraph()
        run = p.add_run(f"[图片: {alt_text}]")
        run.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)
        return False
    return None

def process_markdown_file(doc, md_path, images_dir):
    """处理单个Markdown文件"""
    print(f"正在处理: {md_path}")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            add_heading_custom(doc, line[2:], 1)
            i += 1
            continue
        elif line.startswith('## '):
            add_heading_custom(doc, line[3:], 2)
            i += 1
            continue
        elif line.startswith('### '):
            add_heading_custom(doc, line[4:], 3)
            i += 1
            continue
        
        # 处理图片
        if line.strip().startswith('!['):
            result = process_image_reference(doc, line, images_dir)
            if result is not None:
                i += 1
                continue
        
        # 处理分隔线
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # 处理代码块
        if line.strip().startswith('```'):
            i = process_code_block(doc, lines, i)
            continue
        
        # 处理表格
        if line.strip().startswith('|') and '|' in line[1:]:
            i = process_markdown_table(doc, lines, i)
            continue
        
        # 处理引用块
        if line.strip().startswith('> '):
            quote_text = line[2:]
            # 收集多行引用
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                quote_text += '\n' + lines[i][2:]
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue
        
        # 处理列表
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.+)$', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            content_text = list_match.group(3)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            
            if list_match.group(2) in ['-', '*']:
                p.style = 'List Bullet'
            else:
                p.style = 'List Number'
            
            process_inline_formatting(p, content_text)
            i += 1
            continue
        
        # 普通段落
        p = doc.add_paragraph()
        process_inline_formatting(p, line)
        i += 1

def create_cover(doc, title_text, subtitle_text, date_text):
    """创建封面"""
    # 添加封面
    title = doc.add_paragraph()
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = 'Microsoft YaHei'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(subtitle_text)
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.name = 'Microsoft YaHei'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date_text)
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 分页
    doc.add_page_break()

def main():
    """主函数"""
    import argparse
    from datetime import datetime
    
    parser = argparse.ArgumentParser(description='Markdown to Word 转换器')
    parser.add_argument('--input', '-i', default='chapters', help='章节目录')
    parser.add_argument('--images', '-img', default='images', help='图片目录')
    parser.add_argument('--output', '-o', default=None, help='输出文件')
    parser.add_argument('--title', '-t', default='研究报告', help='报告标题')
    parser.add_argument('--subtitle', '-s', default='', help='副标题')
    
    args = parser.parse_args()
    
    # 获取工作目录（当前目录）
    work_dir = Path.cwd()
    chapters_dir = work_dir / args.input
    images_dir = work_dir / args.images
    
    if args.output:
        output_path = work_dir / args.output
    else:
        output_path = work_dir / f"{args.title}.docx"
    
    print("=" * 60)
    print("Markdown to Word 转换器")
    print("=" * 60)
    
    # 创建文档
    doc = create_document()
    
    # 创建封面
    current_date = datetime.now().strftime('%Y年%m月')
    create_cover(doc, args.title, args.subtitle, current_date)
    
    # 添加目录 - 动态生成
    add_heading_custom(doc, '目录', 1)
    
    # 从章节文件提取标题
    toc_items = []
    for chapter_path in sorted(chapters_dir.glob('*.md')):
        content = chapter_path.read_text(encoding='utf-8')
        # 提取第一个#标题
        match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if match:
            toc_items.append(match.group(1))
        else:
            toc_items.append(chapter_path.stem)
    
    # 添加参考文献到目录
    refs_path = work_dir / 'references' / '参考文献.md'
    if refs_path.exists():
        toc_items.append('参考文献')
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # 处理各章节 - 自动发现chapters目录下的所有MD文件
    chapter_files = sorted([f for f in chapters_dir.glob('*.md') if f.is_file()])
    
    if not chapter_files:
        print(f"  警告: 在 {chapters_dir} 中未找到Markdown文件")
    else:
        print(f"  发现 {len(chapter_files)} 个章节文件")
        for chapter_path in chapter_files:
            print(f"正在处理: {chapter_path.name}")
            process_markdown_file(doc, chapter_path, images_dir)
            doc.add_page_break()
    
    # 添加参考文献
    if refs_path.exists():
        print(f"正在处理: 参考文献")
        process_markdown_file(doc, refs_path, images_dir)
    
    # 保存文档
    doc.save(output_path)
    print("=" * 60)
    print(f"转换完成！输出文件: {output_path}")
    print("=" * 60)

if __name__ == '__main__':
    main()
